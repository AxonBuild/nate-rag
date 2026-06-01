"""Document loader supporting PDF and DOCX files."""
import tempfile
import os
import re
from pathlib import Path
from src.ingestion.models import Document


async def load_document(file_bytes: bytes, filename: str) -> Document:
    """Route to correct loader based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return await _load_pdf(file_bytes, filename)
    elif ext == ".docx":
        return await _load_docx(file_bytes, filename)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


async def _load_pdf(pdf_bytes: bytes, filename: str) -> Document:
    """Load PDF using pdfplumber, combining all pages into one document."""
    import pdfplumber

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(pdf_bytes)
        tmp_path = tmp.name

    try:
        page_texts = []
        with pdfplumber.open(tmp_path) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text(layout=True, x_tolerance=2, y_tolerance=2)
                if text:
                    page_texts.append(text)

        combined = "\n\n".join(page_texts)
        combined = _normalize(combined)

        return Document(
            text=combined,
            metadata={"source": filename, "document_name": filename, "total_pages": total_pages},
        )
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


async def _load_docx(docx_bytes: bytes, filename: str) -> Document:
    """Load DOCX using python-docx, preserving paragraph structure."""
    import docx as docx_lib

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    try:
        doc = docx_lib.Document(tmp_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        combined = "\n\n".join(paragraphs)
        combined = _normalize(combined)

        return Document(
            text=combined,
            metadata={"source": filename, "document_name": filename, "total_pages": 1},
        )
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def _normalize(text: str) -> str:
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = '\n'.join(line.rstrip() for line in text.split('\n'))
    return text.strip()
