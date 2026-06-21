"""Load PDF and DOCX files into plain text."""
import os
import re
import tempfile
from pathlib import Path
import fitz
from backend.app.services.ingestion.document.models import Document


async def load_document(file_bytes: bytes, filename: str) -> Document:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return await _load_pdf(file_bytes, filename)
    if ext == ".docx":
        return await _load_docx(file_bytes, filename)
    raise ValueError(f"Unsupported file type: {ext}. Use .pdf or .docx")


async def _load_pdf(pdf_bytes: bytes, filename: str) -> Document:
    # PyMuPDF (fitz) extracts text in a small fraction of pdfplumber's memory — pdfplumber's
    # extract_text() peaks at ~1 GB on a 300-page file (OOM on a 1 GB box), fitz at ~70 MB.

    page_texts: list[str] = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf:
        total_pages = pdf.page_count
        for page in pdf:
            text = page.get_text()
            if text:
                page_texts.append(text)

    combined = _normalize("\n\n".join(page_texts))
    return Document(
        text=combined,
        metadata={
            "source": filename,
            "document_name": filename,
            "total_pages": total_pages,
        },
    )


async def _load_docx(docx_bytes: bytes, filename: str) -> Document:
    import docx as docx_lib

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    try:
        doc = docx_lib.Document(tmp_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        combined = _normalize("\n\n".join(paragraphs))
        return Document(
            text=combined,
            metadata={
                "source": filename,
                "document_name": filename,
                "total_pages": 1,
            },
        )
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def _normalize(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()
